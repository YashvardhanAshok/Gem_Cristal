['id', 'date_of_search', 'tender_id', 'element_put', 'item_description', 'qty', 'start_date', 'end_date', 'end_time', 'day_left_formula', 'emd_amount', 'tender_value', 'item_category', 'consignee_reporting', 'address', 'MSE', 'ministry', 'department', 'branch', 'link_href', 'file_path', 'matches', 'matched_products', 'status', 'L_Placeholder', 'Live']

import pyodbc
import pandas as pd
import pandas as pd
from docx import Document
from datetime import datetime

conn = pyodbc.connect(
    "DRIVER={ODBC Driver 17 for SQL Server};"
    "SERVER=localhost\\SQLEXPRESS;"
    "DATABASE=gem_tenders;"
    "Trusted_Connection=yes;"
)

query = "SELECT * FROM tender_data"
df = pd.read_sql(query, conn)

def termenal(date_of_today):
    total_entry=len(df)

    df['date_of_search'] = pd.to_datetime(df['date_of_search'], errors='coerce').dt.date
    today = pd.to_datetime(date_of_today).date()

    df_today = df[df['date_of_search'] == today]

    if not df_today.empty:
        department_counts = df_today.groupby('department').size().reset_index(name='entry_count')
        total_today = 0
        for _, row in department_counts.iterrows():
            total_today = total_today + int(row['entry_count'])

        print(f"\nEntries for {today} grouped by department:\n")
        print(department_counts.to_string(index=False))
    else:
        print(f"\nNo entries found for {today}.")

    df['Live'] = df['Live'].astype(str).str.strip().str.lower()
    total_live_yes = (df['Live'] == 'yes').sum()
    total_live_no = (df['Live'] == 'no').sum()

    print(f"\nLive Tender: {total_live_yes}")
    print(f"Close Tender: {total_live_no}")
    print(f"Total entries in database: {total_entry}")
    print(f"\ntoday entry:{total_today}\n")

def word(date_of_today):
    df['date_of_search'] = pd.to_datetime(df['date_of_search'], errors='coerce').dt.date
    df['Live'] = df['Live'].astype(str).str.strip().str.lower()
    today = pd.to_datetime(date_of_today).date()

    total_entries = len(df)
    total_live_yes = (df['Live'] == 'yes').sum()
    total_live_no = (df['Live'] == 'no').sum()

    df_today = df[df['date_of_search'] == today]
    department_counts = df_today.groupby('department').size().reset_index(name='entry_count') if not df_today.empty else None

    total_today = 0
    for _, row in department_counts.iterrows():
        total_today = total_today + int(row['entry_count'])

    doc = Document()
    doc.add_heading('Tender Entry Summary Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total entries in database: {total_entries}")
    # doc.add_paragraph(f"New Tender add: {total_today}")
    
    Live_table = doc.add_table(rows=1, cols=2)
    Live_table.style = 'Table Grid'
    Live_hdr_cells = Live_table.rows[0].cells
    Live_hdr_cells[0].text = 'Live Tender'
    Live_hdr_cells[1].text = 'Close Tender'

    live_row_cells = Live_table.add_row().cells
    live_row_cells[0].text = str(total_live_yes)
    live_row_cells[1].text = str(total_live_no)

    doc.add_paragraph(f"New Tender add: {total_today}")

    if department_counts is not None:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Department'
        hdr_cells[1].text = 'Entry Count'
        for _, row in department_counts.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['department']) if row['department'] else 'MINISTRY OF COMMUNICATIONS'
            row_cells[1].text = str(row['entry_count'])
    else:
        doc.add_paragraph("No entries found for today.")

    output_file = "tender_summary_report.docx"
    doc.save(output_file)
    print(f"\n📄 Report saved to: {output_file}")

termenal('2025-05-28')
word('2025-05-28')

# just completed Indian army and air force   

# Total Tender add today: 2,787
# Total tender: 7813

# monday
# 3603, 4987, 4249

# tuseday
# 4743, 4765, 5026

# wednesday
# 6159, 6218, 6754 , 7813

# friday
# 7908, 9288

# satady
# 9514, 10864, 10984

# mun - 1,845
# 11292, 11359
 
# Tus 27
# 11433, 11464, 11468

# Wen - 347
# 11780, 11882, 1579, 1769